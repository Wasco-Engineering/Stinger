#!/usr/bin/env python3
"""Ensure QF87 template exists in repo (copy from I: or create placeholder)."""

from __future__ import annotations

import shutil
import sys
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parents[1]
DEST = PROJECT_ROOT / 'deploy' / 'templates' / 'qf87' / 'QF87_Stinger_TestStand.docx'
SOURCE = Path(
    r'I:\Level 5 Documentation\Quality Forms\QF87 Calibration Certificate_Teststands_Rev 000.docx',
)


def _create_placeholder_template(path: Path) -> None:
    from docx import Document

    doc = Document()
    doc.add_heading('QF87 Stinger Test Stand Calibration Certificate', level=0)
    doc.add_paragraph('Technician ID: {{TECHNICIAN_ID}}')
    doc.add_paragraph('Asset ID: {{ASSET_ID}}')
    doc.add_paragraph('Equipment ID: {{EQUIPMENT_ID}}')
    doc.add_paragraph('Profile: {{PROFILE_LABEL}}')
    doc.add_paragraph('Started: {{STARTED_AT}}')
    doc.add_paragraph('Completed: {{COMPLETED_AT}}')
    doc.add_paragraph('Overall result: {{OVERALL_RESULT}}')
    doc.add_paragraph('')
    doc.add_paragraph('Left port (port_a): {{PORT_A_RESULT}}')
    doc.add_paragraph('{{PORT_A_DETAIL}}')
    doc.add_paragraph('')
    doc.add_paragraph('Right port (port_b): {{PORT_B_RESULT}}')
    doc.add_paragraph('{{PORT_B_DETAIL}}')
    doc.add_paragraph('')
    doc.add_paragraph('Config changes applied:')
    doc.add_paragraph('{{CONFIG_CHANGES}}')
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    print(f'Created placeholder template: {path}', flush=True)


def main() -> int:
    if DEST.exists():
        print(f'Template already exists: {DEST}', flush=True)
        return 0
    DEST.parent.mkdir(parents=True, exist_ok=True)
    if SOURCE.is_file():
        shutil.copy2(SOURCE, DEST)
        print(f'Copied template from {SOURCE}', flush=True)
        return 0
    try:
        _create_placeholder_template(DEST)
    except ImportError:
        print('Install python-docx: pip install python-docx', flush=True)
        return 1
    return 0


if __name__ == '__main__':
    raise SystemExit(main())
