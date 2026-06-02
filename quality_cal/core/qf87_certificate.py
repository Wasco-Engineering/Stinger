"""QF87 Word certificate fill and export to Desktop."""

from __future__ import annotations

import logging
import re
import sys
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Optional

import yaml

from quality_cal.config import QualitySettings
from quality_cal.session import PORT_LABELS, QualityCalibrationSession

logger = logging.getLogger(__name__)

_PLACEHOLDER_RE = re.compile(r'\{\{([A-Z0-9_]+)\}\}')


def resolve_template_path(configured: Path) -> Path:
    if configured.is_file():
        return configured
    if getattr(sys, 'frozen', False):
        bundled = Path(sys._MEIPASS) / 'deploy' / 'templates' / 'qf87' / 'QF87_Stinger_TestStand.docx'
        if bundled.is_file():
            return bundled
    repo = Path(__file__).resolve().parents[2] / 'deploy' / 'templates' / 'qf87' / 'QF87_Stinger_TestStand.docx'
    if repo.is_file():
        return repo
    return configured


def expand_path_template(raw: str) -> Path:
    expanded = raw.replace('%USERPROFILE%', str(Path.home()))
    return Path(expanded).expanduser()


def default_desktop_output_dir() -> Path:
    return Path.home() / 'Desktop' / 'Stinger' / 'CalibrationReports'


def build_certificate_context(
    session: QualityCalibrationSession,
    settings: QualitySettings,
    *,
    equipment_id: str = 'STINGER',
) -> Dict[str, str]:
    started = session.started_at.strftime('%Y-%m-%d %H:%M') if session.started_at else 'N/A'
    completed = session.completed_at.strftime('%Y-%m-%d %H:%M') if session.completed_at else 'N/A'
    overall = 'PASS' if session.overall_passed else 'FAIL'

    def _port_block(port_id: str) -> tuple[str, str]:
        port = session.port_result(port_id)
        label = PORT_LABELS.get(port_id, port_id)
        if not port.points:
            return f'{label}: not run', ''
        result = 'PASS' if port.overall_passed else 'FAIL'
        head = f'{label}: {result}'
        lines = [f'Points measured: {len(port.points)}']
        fit = port.fit_summary
        if fit is not None:
            if fit.transducer_p99_abs_torr is not None:
                lines.append(f'Transducer p99: {fit.transducer_p99_abs_torr:.3f} Torr')
            if fit.alicat_p99_abs_torr is not None:
                lines.append(f'Alicat p99: {fit.alicat_p99_abs_torr:.3f} Torr')
            lines.append(
                'Models applied: yes' if fit.applied_to_stinger_config else 'Models applied: no',
            )
        return head, '\n'.join(lines)

    port_a_result, port_a_detail = _port_block('port_a')
    port_b_result, port_b_detail = _port_block('port_b')

    config_lines: list[str] = []
    for port_id in ('port_a', 'port_b'):
        fit = session.port_result(port_id).fit_summary
        if fit is None or not fit.applied_to_stinger_config:
            continue
        label = PORT_LABELS.get(port_id, port_id)
        if fit.transducer_error_model:
            config_lines.append(f'{label}: transducer_error_model updated')
        if fit.alicat_error_model:
            config_lines.append(f'{label}: alicat_error_model updated')
    if not config_lines:
        config_lines.append('No error models applied this session.')

    return {
        'TECHNICIAN_ID': session.technician_name or '—',
        'ASSET_ID': session.asset_id or '—',
        'EQUIPMENT_ID': equipment_id,
        'PROFILE_LABEL': settings.profile_label,
        'STARTED_AT': started,
        'COMPLETED_AT': completed,
        'OVERALL_RESULT': overall,
        'PORT_A_RESULT': port_a_result,
        'PORT_A_DETAIL': port_a_detail,
        'PORT_B_RESULT': port_b_result,
        'PORT_B_DETAIL': port_b_detail,
        'CONFIG_CHANGES': '\n'.join(config_lines),
    }


def _replace_in_paragraph(paragraph: Any, context: Dict[str, str]) -> None:
    if not paragraph.text:
        return
    text = paragraph.text
    for match in _PLACEHOLDER_RE.finditer(text):
        token = match.group(1)
        if token in context:
            text = text.replace(match.group(0), context[token])
    if text != paragraph.text:
        paragraph.text = text


def fill_qf87_docx(template_path: Path, context: Dict[str, str], output_path: Path) -> Path:
    from docx import Document

    doc = Document(str(template_path))
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, context)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, context)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def export_docx_to_pdf(docx_path: Path, pdf_path: Path) -> bool:
    """Convert DOCX to PDF using Microsoft Word COM if available."""
    try:
        import win32com.client  # type: ignore[import-untyped]
    except ImportError:
        logger.warning('pywin32 not installed; skipping Word PDF export')
        return False

    word = None
    doc = None
    try:
        word = win32com.client.Dispatch('Word.Application')
        word.Visible = False
        doc = word.Documents.Open(str(docx_path.resolve()))
        doc.ExportAsFixedFormat(str(pdf_path.resolve()), ExportFormat=17)
        return pdf_path.is_file()
    except Exception as exc:
        logger.warning('Word PDF export failed: %s', exc)
        return False
    finally:
        if doc is not None:
            doc.Close(False)
        if word is not None:
            word.Quit()


def certificate_base_filename(session: QualityCalibrationSession, settings: QualitySettings) -> str:
    tech = (session.technician_name or 'tech').replace(' ', '_')
    asset = (session.asset_id or 'asset').replace(' ', '_')
    stamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    return f'{settings.report_filename_prefix}_{tech}_{asset}_{stamp}'


def export_certificate_bundle(
    session: QualityCalibrationSession,
    settings: QualitySettings,
    *,
    equipment_id: str = 'STINGER',
    desktop_dir: Path | None = None,
) -> Dict[str, Optional[Path]]:
    """Write QF87 DOCX (+ PDF when Word available) to desktop output dir."""
    from quality_cal.core.report_generator import export_report_pdf

    out_dir = desktop_dir or settings.desktop_output_dir
    out_dir.mkdir(parents=True, exist_ok=True)
    base = certificate_base_filename(session, settings)
    context = build_certificate_context(session, settings, equipment_id=equipment_id)

    template = resolve_template_path(settings.report_template_path)
    if not template.is_file():
        raise FileNotFoundError(f'QF87 template not found: {template}')

    docx_path = out_dir / f'{base}.docx'
    fill_qf87_docx(template, context, docx_path)

    pdf_path = out_dir / f'{base}.pdf'
    pdf_ok = export_docx_to_pdf(docx_path, pdf_path)
    if not pdf_ok:
        pdf_path = export_report_pdf(session, settings, pdf_path)

    records_docx: Optional[Path] = None
    records_pdf: Optional[Path] = None
    if settings.also_write_records_path:
        records_dir = settings.report_output_dir
        records_dir.mkdir(parents=True, exist_ok=True)
        records_docx = records_dir / docx_path.name
        fill_qf87_docx(template, context, records_docx)
        records_pdf = records_dir / pdf_path.name
        if pdf_ok:
            export_docx_to_pdf(records_docx, records_pdf)
        else:
            export_report_pdf(session, settings, records_pdf)

    return {
        'docx': docx_path,
        'pdf': pdf_path,
        'records_docx': records_docx,
        'records_pdf': records_pdf,
    }


def load_equipment_id_from_stinger_config() -> str:
    try:
        from app.core.paths import get_stinger_config_path

        path = get_stinger_config_path()
        if not path.is_file():
            return 'STINGER'
        data = yaml.safe_load(path.read_text(encoding='utf-8'))
        if isinstance(data, dict):
            tp = data.get('test_parameters', {})
            if isinstance(tp, dict):
                return str(tp.get('equipment_id', 'STINGER'))
    except Exception:
        pass
    return 'STINGER'
